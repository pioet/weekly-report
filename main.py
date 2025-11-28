import os
from datetime import date

from flask import (
    Flask, render_template, request, redirect,
    url_for, send_file, flash
)
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy import func
from pathlib import Path
from docx import Document

BASE_DIR = Path(__file__).resolve().parent

app = Flask(__name__)
app.secret_key = "weekly-report"

# SQLite 数据库
app.config["SQLALCHEMY_DATABASE_URI"] = f"sqlite:///{BASE_DIR / 'weekly_reports.db'}"
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False

# 模板文件存放路径
UPLOAD_DIR = BASE_DIR / "data"
UPLOAD_DIR.mkdir(exist_ok=True)

db = SQLAlchemy(app)


class WeeklyReport(db.Model):
    __tablename__ = "weekly_reports"

    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(64), nullable=True)
    date = db.Column(db.Date, nullable=False)
    summary = db.Column(db.Text, nullable=True)
    plan = db.Column(db.Text, nullable=True)
    created_at = db.Column(db.DateTime, server_default=func.now())
    updated_at = db.Column(db.DateTime, onupdate=func.now(), server_default=func.now())


class ExportConfig(db.Model):
    __tablename__ = "export_config"

    id = db.Column(db.Integer, primary_key=True)
    filename_pattern = db.Column(db.String(255), nullable=False, default="{DATE}_{NAME}_周报.docx")


def init_db():
    db.create_all()
    # 保证有一条 ExportConfig 记录
    if not db.session.get(ExportConfig, 1):
        cfg = ExportConfig(id=1, filename_pattern="{DATE}_{NAME}_周报.docx")
        db.session.add(cfg)
        db.session.commit()



def get_last_name():
    """获取最近一条周报的姓名，用作新建默认值"""
    last_report = WeeklyReport.query.order_by(WeeklyReport.id.desc()).first()
    return last_report.name if last_report and last_report.name else ""


def get_template_path():
    """固定模板路径"""
    return UPLOAD_DIR / "weekly_report_template.docx"


def has_template():
    return get_template_path().exists()


def render_docx_template(template_path, context, output_path):
    """
    用 python-docx 做占位符替换：
    替换 {{NAME}}, {{DATE}}, {{SUMMARY}}, {{PLAN}}

    方案说明：
    - 不再按 run 替换（Word 容易把占位符拆成多个 run，导致找不到完整字符串）
    - 改为按 paragraph.text / cell.paragraph.text 级别处理，
      用拼接后的整段文本做 replace，然后再回写回去。
    - 会丢失这一段内部更细粒度的格式（比如一个段落里既有加粗又有普通字），
      但占位符文本本身通常是整块纯文本，影响不大。
    """
    doc = Document(template_path)

    def replace_in_paragraph(paragraph):
        if not paragraph.text:
            return
        text = paragraph.text
        new_text = text
        for key, value in context.items():
            placeholder = "{{" + key + "}}"
            if placeholder in new_text:
                new_text = new_text.replace(placeholder, value)
        if new_text != text:
            paragraph.text = new_text

    # 1. 普通段落
    for p in doc.paragraphs:
        replace_in_paragraph(p)

    # 2. 表格中的单元格段落
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p)

    # 3. 如果模板里在页眉/页脚中也用了占位符，可以顺便处理
    for section in doc.sections:
        header = section.header
        for p in header.paragraphs:
            replace_in_paragraph(p)
        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_in_paragraph(p)

        footer = section.footer
        for p in footer.paragraphs:
            replace_in_paragraph(p)
        for table in footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_in_paragraph(p)

    doc.save(output_path)



@app.route("/")
def index():
    reports = WeeklyReport.query.order_by(WeeklyReport.date.desc(), WeeklyReport.id.desc()).all()
    return render_template("list.html", reports=reports)


@app.route("/report/new", methods=["GET", "POST"])
def new_report():
    if request.method == "POST":
        name = request.form.get("name") or ""
        date_str = request.form.get("date")
        summary = request.form.get("summary") or ""
        plan = request.form.get("plan") or ""

        # 日期解析（简单处理，假设格式 yyyy-mm-dd）
        try:
            y, m, d = [int(x) for x in date_str.split("-")]
            d_obj = date(y, m, d)
        except Exception:
            flash("日期格式错误，应为 YYYY-MM-DD", "error")
            return redirect(url_for("new_report"))

        report = WeeklyReport(
            name=name.strip(),
            date=d_obj,
            summary=summary,
            plan=plan,
        )
        db.session.add(report)
        db.session.commit()
        flash("周报已创建", "success")
        return redirect(url_for("index"))

    default_name = get_last_name()
    today = date.today().isoformat()
    return render_template("form.html", report=None, default_name=default_name, default_date=today)


@app.route("/report/<int:report_id>/edit", methods=["GET", "POST"])
def edit_report(report_id):
    report = WeeklyReport.query.get_or_404(report_id)

    if request.method == "POST":
        report.name = (request.form.get("name") or "").strip()
        date_str = request.form.get("date")
        report.summary = request.form.get("summary") or ""
        report.plan = request.form.get("plan") or ""

        try:
            y, m, d = [int(x) for x in date_str.split("-")]
            report.date = date(y, m, d)
        except Exception:
            flash("日期格式错误，应为 YYYY-MM-DD", "error")
            return redirect(url_for("edit_report", report_id=report_id))

        db.session.commit()
        flash("周报已保存", "success")
        return redirect(url_for("index"))

    return render_template("form.html", report=report, default_name=report.name, default_date=report.date.isoformat())


@app.route("/settings", methods=["GET", "POST"])
def settings():
    cfg = db.session.get(ExportConfig, 1)
    if request.method == "POST":
        # 文件命名模式
        filename_pattern = request.form.get("filename_pattern") or "{DATE}_{NAME}_周报.docx"
        cfg.filename_pattern = filename_pattern.strip()

        # 处理模板上传
        f = request.files.get("template_file")
        if f and f.filename:
            # 这里只接受 .docx
            if not f.filename.lower().endswith(".docx"):
                flash("只支持 .docx 模板文件", "error")
                return redirect(url_for("settings"))

            template_path = get_template_path()
            f.save(template_path)
            flash("模板和命名格式已保存", "success")
        else:
            flash("命名格式已保存（模板未更改）", "success")

        db.session.commit()
        return redirect(url_for("settings"))

    template_exists = has_template()
    return render_template("settings.html", config=cfg, template_exists=template_exists)


@app.route("/report/<int:report_id>/export")
def export_report(report_id):
    report = WeeklyReport.query.get_or_404(report_id)

    if not has_template():
        flash("尚未上传 Word 模板，请先到设置页面上传模板。", "error")
        return redirect(url_for("settings"))

    cfg = db.session.get(ExportConfig, 1)

    context = {
        "NAME": report.name or "",
        "DATE": report.date.isoformat(),
        "SUMMARY": report.summary or "",
        "PLAN": report.plan or "",
    }

    # 输出文件名
    safe_name = report.name or "未命名"
    filename = cfg.filename_pattern.format(
        NAME=safe_name,
        DATE=report.date.isoformat()
    )

    # 输出路径
    output_path = UPLOAD_DIR / filename
    render_docx_template(get_template_path(), context, output_path)

    return send_file(
        output_path,
        as_attachment=True,
        download_name=filename
    )
    
@app.route("/report/<int:report_id>/delete", methods=["POST"])
def delete_report(report_id):
    report = WeeklyReport.query.get_or_404(report_id)
    db.session.delete(report)
    db.session.commit()
    flash("周报已删除", "success")
    return redirect(url_for("index"))

if __name__ == "__main__":
    with app.app_context():
        init_db()
    app.run(debug=True)
